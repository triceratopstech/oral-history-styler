import csv
import sys
import docx
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import Inches, Pt
import re
from datetime import timedelta

def transcribe(rawCSV):

    #configure this to adjust timestamps.
    #Smithsonian oral history spec is 5 minutes
    timestampIntervalMinutes = 5 #in minutes


    csvreader = csv.reader(rawCSV)
    transcript = str() #start a string to append to
    next(csvreader) #assume first line is just headers and ignore
    timestampInterval = timedelta(minutes=timestampIntervalMinutes)
    nextTimestampMinute = timestampInterval
    speaker_name = re.compile(r"^[A-Za-z\-\"']+$")


    for row in csvreader:
        nextEntry = str()
        try:
            timeString = row[0][0:8]
            timeString = timeString.replace(";",":")
            currentTimestamp = convert_to_timedelta(timeString)

            #is it time to set a timestamp?
            if(currentTimestamp >= nextTimestampMinute): 
                nextEntry = nextEntry + "\n\n["+str(currentTimestamp)+"]"
                nextTimestampMinute = nextTimestampMinute+timestampInterval
                if(not speaker_name.match(row[2])):
                    nextEntry = nextEntry + "\n\n"

            #some cells have multiple lines and some switch speakers part way through
            #so treat it like a series of lines even though it's usually only one
            for line in row[2].splitlines():
                if(speaker_name.match(line)):
                    nextEntry = nextEntry + "\n\n" + line + ' '
                else:
                    nextEntry = nextEntry + line + ' '

            transcript = transcript+nextEntry
    
        except Exception as e:
            print("skipping row with malformed text: " + ','.join(row) + '\n', e)

    return transcript

def openUTF8(filename):
    with open(filename, encoding='utf8') as my_file:
        return(transcribe(my_file))

def makedocx(content):
    doc = docx.Document()

    doc = add_begining(doc)
    doc.add_paragraph(content)
    doc = add_ending(doc)
    doc = finalizeDocFormatting(doc)

    return doc

#one inch margins on all sides, 12pt Times New Roman font
#per https://www.aaa.si.edu/oral-history-program-style-guide-section-1-formatting
def finalizeDocFormatting(doc):
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)

    return doc

def add_ending(doc):
    ending_content = "[END OF INTERVIEW.]"
    doc.add_paragraph(ending_content)
    return doc

def add_begining(doc):

    intro_paragraph = doc.add_paragraph()

    heading_run = intro_paragraph.add_run("Transcript\n\nPreface\n\n")
    heading_run.bold = True

    current_run = intro_paragraph.add_run("The following oral history transcript is the result of a recorded interview with ")

    current_run = intro_paragraph.add_run("Interview Name")
    current_run.font.highlight_color = WD_COLOR_INDEX.YELLOW  #Word Doc color index, enum from MSFT.  Yellow is currently 7

    current_run = intro_paragraph.add_run(" on")

    current_run = intro_paragraph.add_run(" Interview Date")
    current_run.font.highlight_color = WD_COLOR_INDEX.YELLOW

    current_run = intro_paragraph.add_run(". The interview took place at")

    current_run = intro_paragraph.add_run(" Interview Location")
    current_run.font.highlight_color = WD_COLOR_INDEX.YELLOW

    current_run = intro_paragraph.add_run(". The interview was conducted in-person. This interview is part of The "
                                      "National Native American Boarding School Healing Coalition’s Oral History Project.\n\n")

    current_run = intro_paragraph.add_run("Interviewee Name and Interviewer Name")
    current_run.font.highlight_color = WD_COLOR_INDEX.YELLOW

    current_run = intro_paragraph.add_run(" have reviewed the transcript. Their corrections and emendations appear below in brackets"
                                      " with initials. This transcript has been lightly edited for readability by the Archives of"
                                      " American Art. The reader should bear in mind that they are reading a transcript of spoken"
                                      ", rather than written, prose.\n\n\n")


    current_run = intro_paragraph.add_run("Interview") #the interview body is another new paragraph so it will make another newline
    current_run.bold = True


    return doc



def convert_to_timedelta(total_time):
    hours = 0
    minutes = 0
    seconds = 0

    hours, minutes, seconds = map(float, total_time.split(':'))
    timedelta(hours=hours, minutes=minutes, seconds=seconds)

    return timedelta(hours=hours, minutes=minutes, seconds=seconds)

#this is still executable on its own with a csv file as an argument
if __name__ == "__main__":
    with open(sys.argv[1], encoding='utf8') as my_file:
        print(transcribe(my_file))


